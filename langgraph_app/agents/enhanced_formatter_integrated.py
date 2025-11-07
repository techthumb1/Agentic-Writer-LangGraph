# langgraph_app/agents/formatter.py

import re
import logging
from typing import Dict, Any, List, Tuple
from datetime import datetime
from pathlib import Path

from langgraph_app.core.state import EnrichedContentState
from langgraph_app.core.types import FormattedContent, AgentType

logger = logging.getLogger(__name__)

class EnhancedFormatterAgent:
    """Enterprise Formatter Agent for multi-platform content preparation."""
    
    def __init__(self):
        self.agent_type = AgentType.FORMATTER
    
    def execute(self, state: EnrichedContentState) -> EnrichedContentState:
        """Format content for multiple output platforms."""
        # Get content source
        if state.edited_content:
            content = state.edited_content.body
            title = state.edited_content.title
        elif state.draft_content:
            content = state.draft_content.body if hasattr(state.draft_content, 'body') else str(state.draft_content)
            title = state.draft_content.title if hasattr(state.draft_content, 'title') else state.content_spec.topic
        else:
            content = state.content
            title = state.content_spec.topic if state.content_spec else "Untitled"
        
        logger.info("Starting enterprise formatting process...")
        
        # Step 1: Clean AI tells and normalize typography
        content = self._remove_ai_tells(content)
        
        # Step 2: Generate platform-specific formats
        formats = self._generate_formats(content, title, state)
        
        # Step 3: Create formatted content object
        state.formatted_content = FormattedContent(
            markdown=formats["markdown"],
            html=formats.get("html"),
            metadata={
                "formats_available": list(formats.keys()),
                "formatted_at": datetime.now().isoformat(),
                "word_count": len(content.split()),
                "platform_ready": ["medium", "substack", "email", "web"]
            }
        )
        
        # Update legacy content field
        state.content = formats["markdown"]
        
        state.log_agent_execution(self.agent_type, {
            "status": "completed",
            "formats_generated": list(formats.keys()),
            "ai_tells_removed": True
        })
        
        logger.info(f"Formatter completed - {len(formats)} formats generated")
        return state
    
    def _remove_ai_tells(self, content: str) -> str:
        """Remove telltale AI writing patterns."""
        # Replace em dashes with regular hyphens
        content = content.replace("—", "-")
        content = content.replace("–", "-")
        
        # Remove common AI hedging phrases
        ai_phrases = [
            r"\bIt'?s important to note that\b",
            r"\bIt'?s worth noting that\b",
            r"\bIt should be noted that\b",
            r"\bIn today'?s fast-paced world\b",
            r"\bIn conclusion,?\b",
            r"\bLast but not least,?\b",
            r"\bAt the end of the day,?\b",
            r"\bdelve into\b",
            r"\bleverage\b(?! )(?=d|s|ing)",  # Keep "leverage" as noun
        ]
        
        for pattern in ai_phrases:
            content = re.sub(pattern, "", content, flags=re.IGNORECASE)
        
        # Normalize multiple spaces
        content = re.sub(r' {2,}', ' ', content)
        
        # Remove space before punctuation
        content = re.sub(r' ([.,;:!?])', r'\1', content)
        
        return content.strip()
    
    def _generate_formats(self, content: str, title: str, state: EnrichedContentState) -> Dict[str, str]:
        """Generate content in multiple formats."""
        formats = {}
        
        # Base markdown (cleaned)
        formats["markdown"] = self._format_markdown(content, title, state)
        
        # HTML for web
        formats["html"] = self._format_html(content, title, state)
        
        # Medium-ready format
        formats["medium"] = self._format_medium(content, title, state)
        
        # Substack-ready format
        formats["substack"] = self._format_substack(content, title, state)
        
        # Email-ready format
        formats["email"] = self._format_email(content, title, state)
        
        # Generate binary formats and save file paths
        try:
            formats["pdf_path"] = self._generate_pdf(content, title, state)
            formats["docx_path"] = self._generate_docx(content, title, state)
            formats["excel_path"] = self._generate_excel(content, title, state)
        except Exception as e:
            logger.warning(f"Binary format generation failed: {e}")
        
        return formats
    
    def _format_markdown(self, content: str, title: str, state: EnrichedContentState) -> str:
        """Format as clean markdown."""
        output = f"# {title}\n\n"
        
        # Add metadata header if SEO analysis exists
        if state.seo_analysis:
            output = f"""---
title: {state.seo_analysis.meta_title}
description: {state.seo_analysis.meta_description}
keywords: {', '.join(list(state.seo_analysis.keyword_density.keys())[:5])}
author: WriterzRoom
date: {datetime.now().strftime('%Y-%m-%d')}
---

"""
            output += f"# {title}\n\n"
        
        output += content
        
        # Add footer
        output += f"\n\n---\n*Generated with WriterzRoom • {datetime.now().strftime('%B %d, %Y')}*"
        
        return output
    
    def _format_html(self, content: str, title: str, state: EnrichedContentState) -> str:
        """Format as clean HTML with proper structure."""
        # Convert markdown to HTML (basic)
        html_content = self._markdown_to_html(content)
        
        meta_description = ""
        if state.seo_analysis:
            meta_description = f'<meta name="description" content="{state.seo_analysis.meta_description}">'
        
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    {meta_description}
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 2rem;
            color: #333;
        }}
        h1 {{ font-size: 2.5rem; margin-bottom: 1rem; }}
        h2 {{ font-size: 2rem; margin-top: 2rem; border-bottom: 2px solid #eee; padding-bottom: 0.5rem; }}
        h3 {{ font-size: 1.5rem; margin-top: 1.5rem; }}
        p {{ margin-bottom: 1rem; }}
        ul, ol {{ margin-bottom: 1rem; padding-left: 2rem; }}
        code {{ background: #f4f4f4; padding: 0.2rem 0.4rem; border-radius: 3px; }}
        pre {{ background: #f4f4f4; padding: 1rem; border-radius: 5px; overflow-x: auto; }}
        blockquote {{ border-left: 4px solid #ddd; margin: 1rem 0; padding-left: 1rem; color: #666; }}
        footer {{ margin-top: 3rem; padding-top: 1rem; border-top: 1px solid #eee; color: #999; font-size: 0.9rem; }}
    </style>
</head>
<body>
    <article>
        <h1>{title}</h1>
        {html_content}
    </article>
    <footer>
        Generated with WriterzRoom • {datetime.now().strftime('%B %d, %Y')}
    </footer>
</body>
</html>"""
        return html
    
    def _format_medium(self, content: str, title: str, state: EnrichedContentState) -> str:
        """Format for Medium publication."""
        # Medium uses markdown, but with specific formatting preferences
        output = f"# {title}\n\n"
        
        # Add subtitle if available
        if state.seo_analysis and state.seo_analysis.meta_description:
            output += f"*{state.seo_analysis.meta_description}*\n\n---\n\n"
        
        # Medium prefers clear section breaks
        content = re.sub(r'\n\n+', '\n\n---\n\n', content, count=2)  # First 2 breaks only
        
        output += content
        
        return output
    
    def _format_substack(self, content: str, title: str, state: EnrichedContentState) -> str:
        """Format for Substack newsletter."""
        output = f"# {title}\n\n"
        
        # Add engaging opener
        if state.seo_analysis and state.seo_analysis.meta_description:
            output += f"*{state.seo_analysis.meta_description}*\n\n"
        
        output += content
        
        # Add CTA footer
        output += "\n\n---\n\n"
        output += "💡 *Found this helpful? Share it with your network!*\n\n"
        output += "[Share on Twitter](https://twitter.com/intent/tweet) • "
        output += "[Share on LinkedIn](https://www.linkedin.com/sharing/share-offsite/)\n"
        
        return output
    
    def _format_email(self, content: str, title: str, state: EnrichedContentState) -> str:
        """Format for email newsletter (HTML)."""
        # Convert to HTML with email-safe styling
        html_content = self._markdown_to_html_simple(content)
        
        email_html = f"""
<div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; color: #333;">
    <h1 style="color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px;">{title}</h1>
    
    {html_content}
    
    <div style="margin-top: 30px; padding-top: 20px; border-top: 1px solid #ddd; font-size: 12px; color: #999;">
        <p>Generated with WriterzRoom • {datetime.now().strftime('%B %d, %Y')}</p>
    </div>
</div>
"""
        return email_html
    
    def _markdown_to_html(self, markdown: str) -> str:
        """Convert markdown to HTML (basic implementation)."""
        html = markdown
        
        # Headers
        html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
        
        # Bold and italic
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
        
        # Lists (basic)
        html = re.sub(r'^\- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
        html = re.sub(r'(<li>.*</li>)', r'<ul>\1</ul>', html, flags=re.DOTALL)
        
        # Paragraphs
        paragraphs = html.split('\n\n')
        html = '\n'.join([f'<p>{p}</p>' if not p.startswith('<') else p for p in paragraphs if p.strip()])
        
        return html
    
    def _markdown_to_html_simple(self, markdown: str) -> str:
        """Convert markdown to simple HTML for email."""
        html = markdown
        
        # Headers with inline styles
        html = re.sub(r'^### (.+)$', r'<h3 style="color: #2c3e50; font-size: 18px; margin-top: 20px;">\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2 style="color: #2c3e50; font-size: 22px; margin-top: 25px; border-bottom: 2px solid #eee; padding-bottom: 10px;">\1</h2>', html, flags=re.MULTILINE)
        
        # Bold
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        
        # Lists
        html = re.sub(r'^\- (.+)$', r'<li style="margin-bottom: 8px;">\1</li>', html, flags=re.MULTILINE)
        html = re.sub(r'(<li.*?>.*?</li>(?:\n<li.*?>.*?</li>)*)', r'<ul style="margin: 15px 0; padding-left: 20px;">\1</ul>', html, flags=re.DOTALL)
        
        # Paragraphs
        paragraphs = html.split('\n\n')
        html = '\n'.join([f'<p style="margin-bottom: 15px; line-height: 1.6;">{p}</p>' if not p.startswith('<') else p for p in paragraphs if p.strip()])
        
        return html
    
    def _generate_pdf(self, content: str, title: str, state: EnrichedContentState) -> str:
        """Generate PDF file and return path."""
        try:
            from weasyprint import HTML
            
            html_content = self._format_html(content, title, state)
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_title = re.sub(r'[^\w\s-]', '', title).replace(' ', '_')[:50]
            filename = f"{safe_title}_{timestamp}.pdf"
            
            output_dir = Path("/tmp/writerzroom_output")
            output_dir.mkdir(exist_ok=True)
            filepath = output_dir / filename
            
            HTML(string=html_content).write_pdf(filepath)
            
            logger.info(f"PDF generated: {filepath}")
            return str(filepath)
            
        except ImportError:
            logger.warning("WeasyPrint not installed - skipping PDF generation")
            return ""
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            return ""
    
    def _generate_docx(self, content: str, title: str, state: EnrichedContentState) -> str:
        """Generate DOCX file and return path."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if state.seo_analysis:
                meta_para = doc.add_paragraph(state.seo_analysis.meta_description)
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                meta_font = meta_para.runs[0].font
                meta_font.italic = True
                meta_font.size = Pt(11)
                meta_font.color.rgb = RGBColor(128, 128, 128)
                doc.add_paragraph()
            
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith('### '):
                    doc.add_heading(line[4:], 3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], 2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], 1)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif re.match(r'^\d+\.', line):
                    doc.add_paragraph(line.split('.', 1)[1].strip(), style='List Number')
                else:
                    para = doc.add_paragraph(line)
                    if '**' in line:
                        para.clear()
                        parts = re.split(r'(\*\*.*?\*\*)', line)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                run = para.add_run(part[2:-2])
                                run.bold = True
                            else:
                                para.add_run(part)
            
            doc.add_paragraph()
            footer = doc.add_paragraph(f"Generated with WriterzRoom • {datetime.now().strftime('%B %d, %Y')}")
            footer_font = footer.runs[0].font
            footer_font.size = Pt(9)
            footer_font.color.rgb = RGBColor(150, 150, 150)
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_title = re.sub(r'[^\w\s-]', '', title).replace(' ', '_')[:50]
            filename = f"{safe_title}_{timestamp}.docx"
            
            output_dir = Path("/tmp/writerzroom_output")
            output_dir.mkdir(exist_ok=True)
            filepath = output_dir / filename
            
            doc.save(filepath)
            
            logger.info(f"DOCX generated: {filepath}")
            return str(filepath)
            
        except ImportError:
            logger.warning("python-docx not installed - skipping DOCX generation")
            return ""
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            return ""
    
    def _generate_excel(self, content: str, title: str, state: EnrichedContentState) -> str:
        """Generate Excel file with content analysis."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment
            
            wb = Workbook()
            
            ws_content = wb.active
            ws_content.title = "Content"
            
            ws_content['A1'] = title
            ws_content['A1'].font = Font(bold=True, size=16)
            ws_content['A1'].alignment = Alignment(wrap_text=True)
            
            row = 3
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    ws_content.cell(row=row, column=1, value=para.strip())
                    ws_content.cell(row=row, column=1).alignment = Alignment(wrap_text=True, vertical='top')
                    row += 1
            
            ws_content.column_dimensions['A'].width = 100
            
            ws_analytics = wb.create_sheet("Analytics")
            ws_analytics['A1'] = "Content Analytics"
            ws_analytics['A1'].font = Font(bold=True, size=14)
            
            row = 3
            ws_analytics.cell(row=row, column=1, value="Metric").font = Font(bold=True)
            ws_analytics.cell(row=row, column=2, value="Value").font = Font(bold=True)
            row += 1
            
            word_count = len(content.split())
            metrics = [
                ("Word Count", word_count),
                ("Character Count", len(content)),
                ("Paragraph Count", len([p for p in paragraphs if p.strip()])),
            ]
            
            if state.seo_analysis:
                metrics.extend([
                    ("Readability Score", state.seo_analysis.readability_score),
                    ("Top Keywords", ", ".join(list(state.seo_analysis.keyword_density.keys())[:5])),
                ])
            
            for metric, value in metrics:
                ws_analytics.cell(row=row, column=1, value=metric)
                ws_analytics.cell(row=row, column=2, value=value)
                row += 1
            
            ws_analytics.column_dimensions['A'].width = 25
            ws_analytics.column_dimensions['B'].width = 50
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_title = re.sub(r'[^\w\s-]', '', title).replace(' ', '_')[:50]
            filename = f"{safe_title}_{timestamp}.xlsx"
            
            output_dir = Path("/tmp/writerzroom_output")
            output_dir.mkdir(exist_ok=True)
            filepath = output_dir / filename
            
            wb.save(filepath)
            
            logger.info(f"Excel generated: {filepath}")
            return str(filepath)
            
        except ImportError:
            logger.warning("openpyxl not installed - skipping Excel generation")
            return ""
        except Exception as e:
            logger.error(f"Excel generation failed: {e}")
            return ""